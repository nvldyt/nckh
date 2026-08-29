# app.py
# ============================================================
# HỖ TRỢ NGHIÊN CỨU KHOA HỌC – EVIDENCE-BASED RAG
# Bản tối ưu cho luận văn Chuyên khoa cấp I – Dược lâm sàng
# Kiến trúc Modular 4 Engines (Tích hợp Reranker & Study Context)
# Bổ sung tính năng Offline: Summarizer (Python) & Ollama Writer
# ============================================================

import io
import os
import re
import time
import json
import gc

from typing import Any, Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
import streamlit as st
import requests
from docx import Document

from chat_data_engine import render_chat_assistant
from chat_writing_engine import render_writing_chat

from table_selection_engine import (
    StudyObjective, CandidateResult, TableSelectionEngine,
    NarrativePlanner, Priority, Presentation,
)
from statistical_engine import (
    validate_dataframe, descriptive_table, numeric_summary,
    crosstab_test, compare_two_groups, binary_logistic_regression,
    create_clinical_groups, generate_baseline_table,
)
from evidence_engine import (
    SourceDocument, EvidenceChunk, extract_pdf,
    search_pubmed, ingest_pubmed_article, 
    add_source_and_chunks
)
from project_storage import save_project, load_project, list_projects, delete_project
from data_engine import auto_clean_data

from citation_engine import CitationEngine
from audit_engine import Audit_generated_text, internal_overlap_Audit
from retrieval_engine import (
    build_bm25_index, build_embedding_index,
    update_embedding_index, retrieve_evidence,
)
from writing_engine import (
    call_gemini, generate_evidence_based,
    BASE_SYSTEM_RULES, MODEL_LITE, DEFAULT_MODEL,
)
from synthesis_engine import build_literature_matrix
from chapter_assembler_engine import assemble_results_and_discussion_chapter
from offline_writing import render_offline_tab

# --- IMPORT 2 MODULE MỚI ---
from summarizer_engine import render_summarizer_tab
from ollama_writer_engine import render_ollama_writer_tab
from medical_stats_engine import generate_table_one

import requests

# === HÀM CỐT LÕI GỌI OLLAMA TỪ COLAB ===
def call_ollama_colab(prompt: str, url: str, model_name: str = "qwen2.5:14b") -> str:
    """Gửi prompt đến máy chủ Ollama trên Colab và nhận về câu trả lời."""
    if not url:
        return "⚠️ Báo lỗi: Vui lòng dán link Cloudflare từ Colab vào thanh bên trước!"
    
    api_endpoint = f"{url.rstrip('/')}/api/generate"
    payload = {
        "model": model_name,
        "prompt": prompt,
        "stream": False
    }
    
    try:
        response = requests.post(api_endpoint, json=payload, timeout=120)
        response.raise_for_status() 
        result = response.json()
        return result.get("response", "")
    except requests.exceptions.RequestException as e:
        return f"❌ Mất kết nối với Colab. Chi tiết lỗi: {e}"

# ============================================================
# 1. CẤU HÌNH
# ============================================================
st.set_page_config(page_title="NCKH", page_icon="🔬", layout="wide")

DEFAULT_TOP_K = 8
MAX_TOP_K = 20
DEFAULT_VN_JOURNAL_DOMAINS = [
    "tapchiyhocvietnam.vn", "vjol.info", "tapchinghiencuuyhoc.vn",
    "jmp.huemed-univ.edu.vn", "jmpm.vn", "huejmp.vn",
    "tcydls108.benhvien108.vn", "tapchiyhcd.vn", "thaibinhjmp.vn", "hup.edu.vn",
]

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
html,body,p,h1,h2,h3,h4,h5,h6,span,div,label,li,.stMarkdown{font-family:'Inter',sans-serif}
.stApp{background:#f8f9fa;color:#2c3e50}
h1{color:#1e293b!important;text-align:center;font-weight:800;margin-top:0!important;margin-bottom:6px}
h2,h3,h4{color:#0f172a!important;font-weight:600}
.stTabs [data-baseweb="tab-panel"]{background:#fff;border-radius:12px;padding:32px;box-shadow:0 1px 3px rgba(0,0,0,.05);border:1px solid #e2e8f0;margin-top:10px}
.stTabs [data-baseweb="tab-list"]{background:#f1f5f9;border-radius:10px;padding:4px;gap:4px}
.stTabs [data-baseweb="tab"]{border-radius:8px!important;font-weight:600;color:#64748b}
.stTabs [aria-selected="true"]{background:#fff!important;color:#2563eb!important;box-shadow:0 1px 3px rgba(0,0,0,.1)}
div.stButton>button,div.stDownloadButton>button{background:#fff!important;color:#334155!important;font-weight:600;border-radius:8px;border:1px solid #cbd5e1!important;padding:8px 16px}
div.stButton>button[kind="primary"]{background:#2563eb!important;color:#fff!important;border:none!important}
.warning-box{border-left:4px solid #f59e0b;padding:12px 16px;background:#fffbeb;border-radius:6px;color:#92400e;font-size:.95rem}
.danger-box{border-left:4px solid #ef4444;padding:12px 16px;background:#fef2f2;border-radius:6px;color:#991b1b;font-size:.95rem}
.success-box{border-left:4px solid #10b981;padding:12px 16px;background:#ecfdf5;border-radius:6px;color:#065f46;font-size:.95rem}
</style>
""", unsafe_allow_html=True)

UI_NAMESPACE = "nckh_cki"


def init_state():
    if "ui_version" not in st.session_state:
        st.session_state["ui_version"] = 0
    defaults = {
        "documents": {}, "chunks": [], "embeddings": None, "bm25": None,
        "citation_registry": {}, "Audit_log": [], "last_generated": "",
        "last_evidence": [], "current_references": [],
        "vn_journal_domains": list(DEFAULT_VN_JOURNAL_DOMAINS),
        "t3_pm_data": [], "t3_vn_data": [], "t3_query": "",
        "t3_en_keyword": "", "t3_vn_keyword": "",
        "result_cart": [], "saved_tables": {},
        "selection_decisions": [], "narrative_plan": {},
        "study_context": {}, "literature_matrix_df": pd.DataFrame(),
        "assembled_ch3": "", "assembled_ch4": "",
        "excel_data": None, "excel_df": None, "clean_logs": [],
        "ai_pending_remark": "",
        "cached_summary": "",
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value.copy() if isinstance(value, (list, dict, pd.DataFrame)) else value


def ui_key(widget_name: str) -> str:
    return f"{UI_NAMESPACE}_v{st.session_state.get('ui_version',0)}_{widget_name}"


def reset_ui_state():
    st.session_state["ui_version"] = st.session_state.get("ui_version",0) + 1


def _field(obj: Any, name: str, default: Any = None) -> Any:
    return obj.get(name, default) if isinstance(obj, dict) else getattr(obj, name, default)


# ============================================================
# 2. WRAPPERS – KẾT NỐI UI VỚI ENGINES
# ============================================================
def rebuild_index(new_chunks: Optional[List[Dict[str, Any]]] = None):
    all_chunks = st.session_state.get("chunks", [])
    if not all_chunks:
        st.session_state["embeddings"] = None
        st.session_state["bm25"] = None
        gc.collect()
        return
    old = st.session_state.get("embeddings")
    try:
        if new_chunks and old is not None:
            st.session_state["embeddings"] = update_embedding_index(new_chunks, old)
        else:
            st.session_state["embeddings"] = build_embedding_index(all_chunks)
        del old
        gc.collect()
        st.session_state["bm25"] = build_bm25_index(all_chunks)
    except Exception:
        gc.collect()
        raise


def retrieve_evidence_wrapper(query: str, k: int = DEFAULT_TOP_K) -> List[Dict[str, Any]]:
    return retrieve_evidence(
        query=query,
        chunks=st.session_state.get("chunks", []),
        matrix=st.session_state.get("embeddings"),
        bm25=st.session_state.get("bm25"),
        top_k=k,
    )


def get_citation_engine() -> CitationEngine:
    if "citation_engine" not in st.session_state:
        st.session_state["citation_engine"] = CitationEngine()
    return st.session_state["citation_engine"]


def source_metadata(source_id: str) -> Dict[str, Any]:
    return st.session_state.get("documents", {}).get(source_id, {})


def citation_bibliography_wrapper() -> str:
    rows = []
    docs = st.session_state.get("documents", {})
    for ref in st.session_state.get("current_references", []):
        rid = ref.get("ref_id", "")
        sid = rid.replace("REF-", "", 1) if rid.startswith("REF-") else rid
        meta = docs.get(sid, ref.get("metadata", {})) or {}
        authors = meta.get("authors") or "Tác giả chưa xác định"
        title = meta.get("title") or meta.get("file_name") or "Tài liệu chưa xác định"
        journal = meta.get("journal") or "Tài liệu lưu trữ"
        year = meta.get("year") or "Năm chưa rõ"
        text = f"[{ref.get('vancouver_index','')}] {authors}. {title}. {journal}. {year}."
        if meta.get("doi"):
            text += f" DOI: {meta['doi']}."
        if meta.get("pmid"):
            text += f" PMID: {meta['pmid']}."
        if meta.get("url") and meta.get("origin") == "Tạp chí VN":
            text += f" [{meta['url']}]"
        rows.append(text)
    return "\n".join(rows)


def generate_evidence_based_wrapper(task: str, query: str, k: int = DEFAULT_TOP_K):
    evidence = retrieve_evidence_wrapper(query, k=k)
    if not evidence:
        return "Tài liệu được cung cấp chưa đủ bằng chứng để kết luận.", [], []
    engine = get_citation_engine()
    final_text, references, invalid_tags = generate_evidence_based(
        task_prompt=task,
        evidence=evidence,
        citation_engine=engine,
        study_context=st.session_state.get("study_context", {}),
    )
    if final_text:
        st.session_state["last_generated"] = final_text
        st.session_state["last_evidence"] = evidence
        st.session_state["current_references"] = references
    return final_text, evidence, invalid_tags


def Audit_generated_text_wrapper(text: str):
    return Audit_generated_text(text, retrieve_evidence_wrapper(text, k=6))


def internal_overlap_Audit_wrapper(text: str, top_k: int = 5):
    return internal_overlap_Audit(text, st.session_state.get("chunks", []), top_k=top_k)


import re
import json

def extract_metadata_from_text_ai_wrapper(text: str) -> dict:
    prompt = f"""
Bạn là chuyên gia trích xuất dữ liệu Y khoa. NHIỆM VỤ: Đọc văn bản thô của bài báo và tìm đúng 5 trường thông tin dưới đây.

QUY TẮC SỐNG CÒN:
1. KHÔNG LẤY TÊN FILE: Tiêu đề bài báo tuyệt đối không được có đuôi .pdf. Hãy tìm tên khoa học thật sự.
2. VƯỢT QUA BÌA QUẢNG CÁO: Nếu thấy chữ "Join ResearchGate", "Downloaded from", hãy lờ đi và đọc tiếp xuống dưới.
3. ĐỊNH DẠNG: Chỉ trả về ĐÚNG MỘT KHỐI JSON. Không nói thêm bất kỳ câu nào khác. Không dùng Markdown (như ```json).

MẪU JSON BẮT BUỘC (Các key phải viết chữ thường):
{{"authors": "...", "title": "...", "year": "...", "journal": "...", "doi": "..."}}

VĂN BẢN CẦN QUÉT:
{text[:15000]}
"""
    try:
        # Ép temperature=0.0 để AI làm việc như một cái máy, không sáng tạo
        res = call_gemini(prompt, model=DEFAULT_MODEL, temperature=0.0)
        if not res: return {}
        
        # Regex bắt sống khối JSON từ dấu { đầu tiên đến dấu } cuối cùng
        import re
        import json
        match = re.search(r'\{.*\}', res, re.DOTALL)
        if match:
            out = json.loads(match.group(0))
            # Chuẩn hóa toàn bộ key về chữ thường một lần nữa cho chắc ăn
            return {str(k).lower(): str(v) for k, v in out.items()}
        return {}
    except Exception:
        return {}

# ============================================================
# 3. HELPER FUNCTIONS
# ============================================================
def parse_vancouver_text_wrapper(pasted_text: str, docs_info: dict) -> dict:
    prompt = f"""
Bạn là chuyên gia dữ liệu. 
Tôi có một danh sách các tài liệu trong hệ thống như sau (Định dạng: MÃ_ID --- Tên file/Tiêu đề tạm):
{json.dumps(docs_info, ensure_ascii=False, indent=2)}

Người dùng vừa dán một danh sách trích dẫn (chuẩn Vancouver) như sau:
{pasted_text}

NHIỆM VỤ:
1. Phân tích các trích dẫn người dùng dán để lấy: authors, title, journal, year, doi.
2. Đối chiếu Tiêu đề/Tác giả với danh sách tài liệu trong hệ thống để tìm ra MÃ_ID phù hợp nhất cho mỗi trích dẫn.
3. TRẢ VỀ DUY NHẤT 1 KHỐI JSON THEO CẤU TRÚC:
{{
    "MÃ_ID_TƯƠNG_ỨNG_1": {{"authors": "...", "title": "...", "year": "...", "journal": "...", "doi": "..."}},
    "MÃ_ID_TƯƠNG_ỨNG_2": {{"authors": "...", "title": "...", "year": "...", "journal": "...", "doi": "..."}}
}}
"""
    try:
        res = call_gemini(prompt, model=DEFAULT_MODEL, temperature=0.0)
        if not res: return {}
        import re
        match = re.search(r'\{.*\}', res, re.DOTALL)
        if match:
            out = json.loads(match.group(0))
            return {str(k): {str(ik).lower(): str(iv) for ik, iv in v.items()} for k, v in out.items() if isinstance(v, dict)}
        return {}
    except Exception:
        return {}

def add_pdf_documents(uploaded_files) -> Tuple[int, int, List[str]]:
    new_sources = 0
    new_chunks_count = 0
    errors = []
    new_chunks_list = []
    for uploaded_file in uploaded_files:
        try:
            source, chunks = extract_pdf(uploaded_file)
            if add_source_and_chunks(source, chunks):
                new_sources += 1
                new_chunks_count += len(chunks)
                new_chunks_list.extend(chunks)
                sid = _field(source, "source_id")
                if sid and chunks:
                    # GOM CHỮ TỪ 3 TRANG ĐẦU (Đề phòng trang bìa ResearchGate)
                    first_chunks = []
                    for c in chunks:
                        page_val = str(_field(c, "page", "")).strip()
                        if page_val in ["1", "2", "3"]:
                            first_chunks.append(_field(c, "text", "") or "")
                    
                    target = "\n".join(first_chunks)
                    if not target:
                        target = "\n".join([_field(c, "text", "") or "" for c in chunks[:10]])
                    
                    target = target[:12000] # Nới rộng giới hạn lên 12000 ký tự
                    
                    if target:
                        meta = extract_metadata_from_text_ai_wrapper(target)
                        if meta and sid in st.session_state.get("documents", {}):
                            for k, v in meta.items():
                                if v and k == "title" and str(v).lower().endswith(".pdf"):
                                    continue
                                if v:
                                    st.session_state["documents"][sid][k] = v
                gc.collect()
        except Exception as exc:
            errors.append(f"{getattr(uploaded_file,'name','PDF')}: {exc}")
            gc.collect()
            
    if new_sources: rebuild_index(new_chunks=new_chunks_list)
    del new_chunks_list
    gc.collect()
    return new_sources, new_chunks_count, errors

def evidence_database_summary() -> Dict[str, Any]:
    docs = st.session_state.get("documents", {})
    chunks = st.session_state.get("chunks", [])
    by_src = {}
    for meta in docs.values():
        o = meta.get("origin", "Khác")
        by_src[o] = by_src.get(o, 0) + 1
    sid_origin = {sid: meta.get("origin", "Khác") for sid, meta in docs.items()}
    by_chunk = {}
    for c in chunks:
        o = sid_origin.get(_field(c, "source_id"), "Khác")
        by_chunk[o] = by_chunk.get(o, 0) + 1
    return {
        "total_sources": len(docs), "total_chunks": len(chunks),
        "by_origin_sources": by_src, "by_origin_chunks": by_chunk,
        "index_ready": st.session_state.get("embeddings") is not None,
    }


def render_evidence_database_status(context_label: str = ""):
    s = evidence_database_summary()
    if not s["total_sources"]:
        st.markdown('<div class="danger-box">⚠️ <b>Evidence Database đang RỖNG.</b> Hãy nạp PDF ở Tab 1 hoặc tra cứu và nạp bài báo ở Tab 2.</div>', unsafe_allow_html=True)
        return
    labels = {"PDF":"📄 PDF","PubMed":"🌍 PubMed","Tạp chí VN":"🇻🇳 Tạp chí VN","Khác":"❓ Khác"}
    pieces = []
    for origin, count in s["by_origin_sources"].items():
        pieces.append(f"{labels.get(origin, origin)}: <b>{count}</b> nguồn ({s['by_origin_chunks'].get(origin,0)} đoạn)")
    st.markdown(
        f'<div class="success-box">✅ <b>Evidence Database{(" - "+context_label) if context_label else ""}:</b> {s["total_sources"]} nguồn / {s["total_chunks"]} đoạn &nbsp;—&nbsp; {" &nbsp;|&nbsp; ".join(pieces)}{("" if s["index_ready"] else " &nbsp;—&nbsp; ⚠️ chưa dựng xong index")}</div>',
        unsafe_allow_html=True,
    )


def translate_query_to_mesh(vietnamese_query: str) -> str:
    prompt = f"""Bạn là một chuyên gia tra cứu tài liệu y khoa.
Chuyển tên đề tài tiếng Việt sau thành chuỗi từ khóa tiếng Anh hiệu quả để tra cứu PubMed.
Dùng AND, OR; không dùng dấu gạch chéo; giữ 2-4 cụm từ khóa; chỉ trả về chuỗi tiếng Anh.
Đề tài: {vietnamese_query}"""
    try:
        text = call_gemini(prompt, model=MODEL_LITE)
    except Exception:
        return vietnamese_query
    return text.strip().strip('"').strip("'").replace("\n", " ") if text else vietnamese_query


def normalize_pubmed_query(en_query: str, fallback_query: str) -> str:
    q = (en_query or "").split("###")[0].split("---")[0].strip()
    markers = ["hoặc", "nếu", "về", "trong", "đề tài", "từ khóa", "từ khoá"]
    if not q or any(x in q.lower() for x in markers):
        return fallback_query.strip()
    return q


def add_markdown_body_to_doc(doc: Document, body: str):
    lines = body.split("\n")
    buffer = []
    def flush():
        if buffer:
            text = " ".join(buffer).strip()
            if text:
                doc.add_paragraph(text)
            buffer.clear()
    for line in lines:
        s = line.strip()
        is_table = s.startswith("|")
        if not is_table:
            doc._last_table_open = False
        if not s:
            flush(); continue
        if s.startswith("### "):
            flush(); doc.add_heading(s[4:].strip(), level=3)
        elif s.startswith("## "):
            flush(); doc.add_heading(s[3:].strip(), level=2)
        elif s.startswith("# "):
            flush(); doc.add_heading(s[2:].strip(), level=1)
        elif s.startswith(("- ", "* ")):
            flush(); doc.add_paragraph(s[2:].strip(), style="List Bullet")
        elif is_table:
            test = set(s.replace("|","").replace(" ","").replace(":",""))
            if test and test <= {"-"}:
                continue
            flush()
            cells = [x.strip() for x in s.strip("|").split("|")]
            table = doc.tables[-1] if doc.tables and getattr(doc,"_last_table_open",False) else None
            if table is None:
                table = doc.add_table(rows=1, cols=max(1,len(cells)))
                table.style = "Light Grid Accent 1"
                for i, cell in enumerate(cells): table.rows[0].cells[i].text = cell
                doc._last_table_open = True
            else:
                row = table.add_row()
                for i, cell in enumerate(cells):
                    if i < len(row.cells): row.cells[i].text = cell
        else:
            buffer.append(s)
    flush()


def create_word_document(title: str, body: str, bibliography: str = "") -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    add_markdown_body_to_doc(doc, body)
    if bibliography.strip():
        doc.add_heading("Tài liệu tham khảo", level=1)
        for item in bibliography.splitlines():
            if item.strip(): doc.add_paragraph(item.strip())
    out = io.BytesIO(); doc.save(out); return out.getvalue()


def safe_read_excel(uploaded_file) -> pd.DataFrame:
    try:
        df = pd.read_excel(uploaded_file)
        df.columns = df.columns.astype(str)
        for col in df.columns:
            if df[col].dtype == "object":
                df[col] = df[col].astype(str)
        gc.collect()
        return df
    except Exception as exc:
        st.error(f"❌ Lỗi khi đọc Excel: {exc}")
        return pd.DataFrame()


def large_batch_ok(name: str, count: int, threshold: int = 250) -> bool:
    if count <= threshold:
        return True
    st.warning(f"⚠️ {name} sẽ xử lý khoảng **{count} tổ hợp** và có thể dùng nhiều RAM.")
    return st.checkbox("Tôi hiểu và muốn chạy toàn bộ tổ hợp này", key=ui_key(f"confirm_{name}_{count}"))


def upsert_candidate(result_id, title, result_type, variables, scientific_value, clinical_importance, discussion_value, p_value=None):
    cart = st.session_state.setdefault("result_cart", [])
    cart[:] = [x for x in cart if getattr(x,"id",None) != result_id]
    cart.append(CandidateResult(
        id=result_id, title=title, result_type=result_type,
        variables=variables, p_value=p_value,
        scientific_value=scientific_value,
        clinical_importance=clinical_importance,
        discussion_value=discussion_value,
    ))


def project_payload() -> dict:
    return {
        "documents": st.session_state.get("documents", {}),
        "chunks": st.session_state.get("chunks", []),
        "citation_registry": st.session_state.get("citation_registry", {}),
        "current_references": st.session_state.get("current_references", []),
        "result_cart": [vars(x) if hasattr(x,"__dict__") else x for x in st.session_state.get("result_cart",[])],
        "saved_tables": {k:v.to_dict(orient="split") for k,v in st.session_state.get("saved_tables",{}).items()},
        "study_context": st.session_state.get("study_context", {}),
        "narrative_plan": st.session_state.get("narrative_plan", {}),
        "assembled_ch3": st.session_state.get("assembled_ch3", ""),
        "assembled_ch4": st.session_state.get("assembled_ch4", ""),
        "vn_journal_domains": st.session_state.get("vn_journal_domains", list(DEFAULT_VN_JOURNAL_DOMAINS)),
        "cached_summary": st.session_state.get("cached_summary", ""),
    }


def restore_project(data: dict):
    st.session_state["documents"] = data.get("documents", {})
    st.session_state["chunks"] = data.get("chunks", [])
    st.session_state["citation_registry"] = data.get("citation_registry", {})
    st.session_state["current_references"] = data.get("current_references", [])
    st.session_state["study_context"] = data.get("study_context", {})
    st.session_state["narrative_plan"] = data.get("narrative_plan", {})
    st.session_state["assembled_ch3"] = data.get("assembled_ch3", "")
    st.session_state["assembled_ch4"] = data.get("assembled_ch4", "")
    st.session_state["vn_journal_domains"] = data.get("vn_journal_domains", list(DEFAULT_VN_JOURNAL_DOMAINS))
    st.session_state["cached_summary"] = data.get("cached_summary", "")
    st.session_state["result_cart"] = []
    for item in data.get("result_cart", []):
        if isinstance(item, dict):
            try: st.session_state["result_cart"].append(CandidateResult(**item))
            except Exception: pass
        else:
            st.session_state["result_cart"].append(item)
    st.session_state["saved_tables"] = {}
    for k,v in data.get("saved_tables",{}).items():
        try: st.session_state["saved_tables"][k] = pd.DataFrame.from_dict(v, orient="split")
        except Exception: pass
    rebuild_index()
    gc.collect()
def format_numbered_citations(generated_text: str) -> str:
    """Hàm chuẩn hóa và đánh số thứ tự trích dẫn nối tiếp chuẩn Vancouver."""
    pattern = r'\[(?:REF-)?(SRC-[A-Z0-9]+)\]'
    citation_mapping = {}
    current_index = 1
    
    def replacer(match):
        nonlocal current_index
        ref_id = match.group(1)
        if ref_id not in citation_mapping:
            citation_mapping[ref_id] = current_index
            current_index += 1
        return f"[{citation_mapping[ref_id]}]"
        
    clean_text = re.sub(pattern, replacer, generated_text)
    
    new_refs = []
    docs = st.session_state.get("documents", {})
    for ref_id, idx in citation_mapping.items():
        new_refs.append({
            "ref_id": ref_id,
            "vancouver_index": idx,
            "metadata": docs.get(ref_id, {})
        })
    
    st.session_state["current_references"] = new_refs
    st.session_state["citation_registry"] = citation_mapping
    
    return clean_text

# ============================================================
# 4. MAIN UI - 9 TABS TÍCH HỢP ĐẦY ĐỦ
# ============================================================
def main():
    init_state()
    st.title("🔬 HỖ TRỢ NGHIÊN CỨU KHOA HỌC")
    st.caption("Evidence-Based RAG • PubMed • Citation Registry • Statistical Engine • Audit")

    with st.sidebar:
        st.header("⚙️ Quản lý dự án")
        if st.button("🧹 Làm mới Giao diện", use_container_width=True):
            reset_ui_state(); st.rerun()
        st.divider(); st.subheader("💾 Lưu & Khôi phục (.json)")
        pdata = json.dumps(project_payload(), ensure_ascii=False, indent=4)
        st.download_button("📥 Tải file dự án", data=pdata, file_name="Du_An_Luan_Van.json", mime="application/json", use_container_width=True)
        uploaded_proj = st.file_uploader("Khôi phục từ file:", type=["json"], key=ui_key("upload_project_json"))
        if uploaded_proj and st.button("🚀 Khôi phục", type="primary", use_container_width=True):
            try:
                restore_project(json.load(uploaded_proj)); st.success("🎉 Khôi phục thành công!"); time.sleep(.8); reset_ui_state(); st.rerun()
            except Exception as exc:
                st.error(f"❌ Lỗi: {exc}")
    
    tabs = st.tabs([
        "🔍 1. PubMed", 
        "📑 2. Tài liệu",
        "🛠️ 3. Tổng hợp", 
        "⚡ 4. Tóm tắt",
        "✍️ 5. Viết văn (Grok AI)", 
        "🤖 6. Viết văn (Ollama)",
        "📊 7. Phân tích số liệu", 
        "🔎 8. Kiểm tra luận văn"
    ])

    # ------------------------------------------------------------
    # TAB 1 - PubMed (ĐÃ TÁCH MODULE)
    # ------------------------------------------------------------
    with tabs[0]:
        from pubmed_tab import render_pubmed_tab
        render_pubmed_tab(
            ui_key=ui_key,
            render_evidence_database_status=render_evidence_database_status,
            translate_query_to_mesh=translate_query_to_mesh,
            search_pubmed=search_pubmed,
            normalize_pubmed_query=normalize_pubmed_query,
            ingest_pubmed_article=ingest_pubmed_article,
            rebuild_index=rebuild_index
        )

    # ------------------------------------------------------------
    # TAB 2 - Tài liệu (PDF) (ĐÃ TÁCH MODULE)
    # ------------------------------------------------------------
    with tabs[1]:
        from pdf_documents_tab import render_pdf_documents_tab
        render_pdf_documents_tab(
            ui_key=ui_key,
            render_evidence_database_status=render_evidence_database_status,
            add_pdf_documents=add_pdf_documents,
            retrieve_evidence_wrapper=retrieve_evidence_wrapper,
            MAX_TOP_K=MAX_TOP_K,
            DEFAULT_TOP_K=DEFAULT_TOP_K,
            extract_metadata_from_text_ai_wrapper=extract_metadata_from_text_ai_wrapper,
            parse_vancouver_text_wrapper=parse_vancouver_text_wrapper,  # <--- THÊM DÒNG NÀY
            _field=_field
        )
    # ------------------------------------------------------------
    # TAB 3 – TỔNG HỢP OFFLINE
    # ------------------------------------------------------------
    with tabs[2]:
        render_offline_tab()

    # ------------------------------------------------------------
    # TAB 4 – TÓM TẮT BẰNG PYTHON
    # ------------------------------------------------------------
    with tabs[3]:
        render_summarizer_tab()

    # ------------------------------------------------------------
    # TAB 5 – VIẾT LUẬN VĂN (GEMINI API - ĐÃ TÁCH MODULE)
    # ------------------------------------------------------------
    with tabs[4]:
        from gemini_writer_tab import render_gemini_writer_tab
        render_gemini_writer_tab(
            ui_key=ui_key,
            render_evidence_database_status=render_evidence_database_status,
            build_literature_matrix=build_literature_matrix,
            create_word_document=create_word_document,
            generate_evidence_based_wrapper=generate_evidence_based_wrapper,
            get_citation_engine=get_citation_engine,
            citation_bibliography_wrapper=citation_bibliography_wrapper,
            Audit_generated_text_wrapper=Audit_generated_text_wrapper,
            internal_overlap_Audit_wrapper=internal_overlap_Audit_wrapper,
            _field=_field,
            format_numbered_citations=format_numbered_citations,
            extract_metadata_from_text_ai_wrapper=extract_metadata_from_text_ai_wrapper
        )
    # ------------------------------------------------------------
    # TAB 6 – VIẾT LUẬN VĂN (OPENROUTER - QWEN 72B)
    # ------------------------------------------------------------
    with tabs[5]:
        from openrouter_writer_engine import render_openrouter_writer_tab
        render_openrouter_writer_tab()
    # ------------------------------------------------------------
    # TAB 7 – PHÂN TÍCH SỐ LIỆU (SPSS MINI - ĐÃ TÁCH MODULE)
    # ------------------------------------------------------------
    with tabs[6]:
        from statistics_tab import render_statistics_tab
        render_statistics_tab(
            ui_key=ui_key,
            safe_read_excel=safe_read_excel,
            auto_clean_data=auto_clean_data,
            render_chat_assistant=render_chat_assistant,
            validate_dataframe=validate_dataframe,
            create_word_document=create_word_document,
            StudyObjective=StudyObjective,
            TableSelectionEngine=TableSelectionEngine,
            NarrativePlanner=NarrativePlanner,
            assemble_results_and_discussion_chapter=assemble_results_and_discussion_chapter,
            format_numbered_citations=format_numbered_citations,
            get_citation_engine=get_citation_engine,
            citation_bibliography_wrapper=citation_bibliography_wrapper,
            descriptive_table=descriptive_table,
            numeric_summary=numeric_summary,
            crosstab_test=crosstab_test,
            large_batch_ok=large_batch_ok,
            compare_two_groups=compare_two_groups,
            binary_logistic_regression=binary_logistic_regression,
            upsert_candidate=upsert_candidate,
            call_gemini=call_gemini,
            BASE_SYSTEM_RULES=BASE_SYSTEM_RULES,
            DEFAULT_MODEL=DEFAULT_MODEL
        )
    # ------------------------------------------------------------
    # TAB 8 – AUDIT (ĐÃ TÁCH MODULE)
    # ------------------------------------------------------------
    with tabs[7]:
        from audit_tab import render_audit_tab
        render_audit_tab(
            ui_key=ui_key,
            Audit_generated_text_wrapper=Audit_generated_text_wrapper,
            internal_overlap_Audit_wrapper=internal_overlap_Audit_wrapper,
            call_gemini=call_gemini,
            BASE_SYSTEM_RULES=BASE_SYSTEM_RULES,
            MODEL_LITE=MODEL_LITE
        )

    
if __name__ == "__main__":
    main()

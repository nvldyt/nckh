# ollama_writer_engine.py
# ============================================================
# MODULE VIẾT LUẬN VĂN BẰNG OLLAMA (RAG TRỰC TIẾP - CHUẨN HỌC THUẬT)
# ============================================================

import streamlit as st
import requests
import io
import re
from docx import Document

try:
    from retrieval_engine import retrieve_evidence
except ImportError:
    retrieve_evidence = None

def call_ollama_colab(prompt: str, url: str, model: str, temperature: float = 0.3) -> str:
    """Gửi prompt đến máy chủ Ollama trên Google Colab qua Cầu nối FastAPI"""
    if not url:
        return "⚠️ Báo lỗi: Vui lòng dán đường dẫn Cloudflare (.trycloudflare.com) từ Google Colab vào thanh bên (Sidebar) trước khi bấm gọi AI!"
    
    url = url.strip()
    md_match = re.search(r'\((https?://[^\s)]+)\)', url)
    if md_match:
        url = md_match.group(1)
    else:
        url = url.strip("[]'\"<>")

    api_endpoint = f"{url.rstrip('/')}/api/generate"
    
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": temperature
        }
    }
    
    try:
        response = requests.post(api_endpoint, json=payload, timeout=300)
        if response.status_code == 200:
            return response.json().get("response", "").strip()
        else:
            return f"❌ Lỗi từ máy chủ Colab (Mã lỗi: {response.status_code})"
    except requests.exceptions.ConnectionError:
        return "❌ Mất kết nối với Colab. Hãy kiểm tra lại đường link Cloudflare."
    except Exception as exc:
        return f"❌ Lỗi kết nối: {exc}"

def create_word_document(title: str, body: str) -> bytes:
    """Tạo file Word (.docx) từ văn bản kết quả"""
    doc = Document()
    doc.add_heading(title, level=0)
    for line in body.split("\n"):
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line.strip()[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            if line.strip(): 
                doc.add_paragraph(line.strip())
    
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def render_ollama_writer_tab():
    """Giao diện Tab 6: Viết luận văn bằng Ollama tích hợp RAG trực tiếp"""
    st.subheader("🤖 Trợ lý Viết luận văn bằng Ollama (RAG Trực tiếp từ Kho tài liệu)")
    st.info("💡 Hệ thống tự động phân tích yêu cầu, nhận biết đang viết phần nào để ép AI tuân thủ đúng chuẩn mực văn phong luận văn y khoa.")

    # --- THANH BÊN (SIDEBAR) ---
    st.sidebar.divider()
    st.sidebar.header("🔌 Kết nối Máy chủ AI (Colab)")
    
    ollama_url = st.sidebar.text_input(
        "API Base URL:", 
        placeholder="https://...trycloudflare.com",
        key="ollama_colab_url_tab6"
    )

    chunks = st.session_state.get("chunks", [])
    if not chunks:
        st.warning("⚠️ Evidence Database đang trống! Hãy nạp tài liệu PDF (Tab 1) hoặc bài báo PubMed (Tab 2) trước.")
        return

    col1, col2 = st.columns(2)
    with col1:
        ollama_model = st.selectbox(
            "Lựa chọn Model trên Colab:", 
            ["qwen2.5:14b", "qwen2.5:7b", "llama3:8b", "qwen2.5:3b"], 
            key="ollama_model_select"
        )
    with col2:
        top_k = st.slider("Số đoạn bằng chứng truy xuất:", 3, 15, 8, key="ollama_top_k")

    user_query = st.text_area(
        "Nhập yêu cầu viết văn hoặc vấn đề cần phân tích học thuật:", 
        placeholder="VD: Viết phần Đặt vấn đề và tính cấp thiết của việc theo dõi nồng độ Vancomycin...", 
        height=120,
        key="ollama_user_query"
    )

    if st.button("🚀 Truy xuất bằng chứng & Yêu cầu Ollama Viết Bài", type="primary", key="btn_run_ollama_writer"):
        if not user_query.strip():
            st.warning("⚠️ Vui lòng nhập nội dung yêu cầu trước khi gửi.")
        else:
            with st.spinner("AI đang xử lý và soạn thảo theo đúng chuẩn học thuật..."):
                
                # 1. Truy xuất bằng chứng
                evidence = []
                if retrieve_evidence:
                    try:
                        evidence = retrieve_evidence(
                            query=user_query,
                            chunks=chunks,
                            matrix=st.session_state.get("embeddings"),
                            bm25=st.session_state.get("bm25"),
                            top_k=top_k
                        )
                    except Exception as e:
                        st.error(f"❌ Lỗi truy xuất: {e}")
                
                evidence_text = "\n".join([f"- {ev.get('text', '')}" for ev in evidence]) if evidence else "Không có bằng chứng cụ thể."

                # 2. KIỂM TRA XEM ANH ĐANG MUỐN VIẾT PHẦN NÀO ĐỂ ÉP PROMPT CHUẨN XÁC
                query_lower = user_query.lower()
                is_intro = "đặt vấn đề" in query_lower or "cấp thiết" in query_lower or "mở đầu" in query_lower

                if is_intro:
                    # PROMPT ĐÃ ĐƯỢC SIẾT CHẶT CỰC KỲ NGHIÊM NGẶT ĐỂ CHỐNG TÓM TẮT VÀ CHỐNG TIẾNG ANH
                    prompt = f"""[LỆNH TỐI CAO: BẮT BUỘC DÙNG TIẾNG VIỆT 100%. CẤM TUYỆT ĐỐI DÙNG TIẾNG ANH. CẤM TÓM TẮT DƯỚI DẠNG GẠCH ĐẦU DÒNG.]

Bạn là một chuyên gia và giảng viên Dược lâm sàng hàng đầu. Nhiệm vụ duy nhất của bạn là VIẾT MỘT BÀI LUẬN HOÀN CHỈNH bằng văn xuôi tiếng Việt cho phần 'ĐẶT VẤN ĐỀ VÀ TÍNH CẤP THIẾT CỦA ĐỀ TÀI' (độ dài khoảng 450-500 từ).

YÊU CẦU BẮT BUỘC:
1. VIẾT THÀNH VĂN XUÔI LIÊN TỤC, chia thành 3 đến 4 đoạn văn học thuật rõ ràng. KHÔNG ĐƯỢC dùng gạch đầu dòng, KHÔNG ĐƯỢC liệt kê dạng key-points.
2. TUYỆT ĐỐI KHÔNG tóm tắt hay sao chép lại cấu trúc của tài liệu tham khảo. Không nhắc đến thời gian nghiên cứu cụ thể hay tiêu chuẩn chọn mẫu thô vào phần này.
3. Nội dung phải tập trung vào: Gánh nặng dịch tễ của nhiễm khuẩn Gram dương/MRSA, thách thức khi sử dụng kháng sinh có khoảng điều trị hẹp như Vancomycin, tầm quan trọng của việc theo dõi nồng độ thuốc (TDM), từ đó nêu bật tính cấp thiết của đề tài.

DỮ LIỆU ĐỂ THAM KHẢO Ý TƯỞNG (KHÔNG ĐƯỢC CHÉP NGUYÊN VĂN):
{evidence_text}
"""
                else:
                    # PROMPT CHUNG CHO CÁC PHẦN KHÁC (Tổng quan, Bàn luận...)
                    prompt = f"""Bạn là một chuyên gia Dược lâm sàng hàng đầu. Viết nội dung học thuật dựa trên yêu cầu và bằng chứng thực tế sau đây.

YÊU CẦU:
{user_query}

BẰNG CHỨNG THỰC TẾ:
{evidence_text}

QUY TẮC: Viết văn xuôi liền mạch, học thuật, sâu sắc về dược động học (PK/PD), không dùng nhãn tóm tắt bài báo, đính kèm mã nguồn [SRC-...] nếu lấy từ dữ liệu thực tế.
"""

                # 3. Gọi Ollama qua Colab
                result = call_ollama_colab(prompt=prompt, url=ollama_url, model=ollama_model)

                # 4. Hiển thị kết quả
                st.markdown("---")
                st.markdown("### 📝 Bản thảo từ Ollama (Chuẩn học thuật):")
                st.markdown(
                    f"<div style='background-color: white; padding: 20px; border-radius: 10px; color: black; border: 1px solid #ddd; font-size: 16px;'>"
                    f"{result}</div>", 
                    unsafe_allow_html=True
                )
                
                with st.expander("🔎 Xem dấu vết bằng chứng (Evidence Trace)"):
                    for ev in evidence:
                        meta = st.session_state.get("documents", {}).get(ev.get("source_id"), {})
                        st.markdown(f"- **Tài liệu:** `{meta.get('file_name', 'N/A')}` | **Mã đoạn:** `{ev.get('chunk_id', '')}`")
                        st.info(f"_{ev.get('text', '')}_")

                st.download_button(
                    label="📥 Tải Bản thảo ra file Word (.docx)",
                    data=create_word_document("Bản thảo nghiên cứu", result),
                    file_name="Ban_thao_Chuan.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_word_ollama_rag"
                )
